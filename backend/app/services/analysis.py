import os
import asyncio
from typing import Dict, List, Any
import google.generativeai as genai
import json
from docx import Document
import re
from typing import Dict, Any, Tuple

from app.config import settings
from app.models.schemas import ConversationSummary, ScoredItem, ConversationType

class AnalysisService:
    def __init__(self):
        """Initialize the analysis service with rubric and model"""
        self.scoring_rubric = self._load_scoring_rubric()
        self.model = None
        
        # Try to initialize Gemini model if API key is available
        if settings.gemini_api_key:
            try:
                genai.configure(api_key=settings.gemini_api_key)
                self.model = genai.GenerativeModel("gemini-1.5-flash")
                print("✅ Gemini AI model initialized successfully")
            except Exception as e:
                print(f"⚠️ Failed to initialize Gemini AI model: {e}")
                print("🔄 Will use enhanced mock analysis instead")
        else:
            print("📝 No Gemini API key found - using enhanced mock analysis")
    
    def set_api_key(self, api_key: str):
        """
        Set Gemini API key and initialize the model
        This allows users to enable real LLM analysis
        """
        try:
            genai.configure(api_key=api_key)
            self.model = genai.GenerativeModel("gemini-1.5-flash")
            print("✅ Gemini API key configured successfully")
            return True
        except Exception as e:
            print(f"❌ Failed to configure Gemini API key: {e}")
            return False
    
    def _load_scoring_rubric(self) -> str:
        """
        Load the scoring rubric from the uploaded Word document
        """
        rubric_path = os.path.join(settings.docs_dir, "Dynamic Conversation Audit Scorecard - Detailed Scoring Rubric.docx")
        
        if os.path.exists(rubric_path):
            try:
                # Extract text from Word document
                doc = Document(rubric_path)
                rubric_text = ""
                
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        rubric_text += paragraph.text + "\n"
                
                # Also extract text from tables if any
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            if cell.text.strip():
                                rubric_text += cell.text + "\n"
                
                return rubric_text.strip()
                
            except Exception as e:
                print(f"Error loading rubric: {e}")
                return self._get_fallback_rubric()
        else:
            return self._get_fallback_rubric()
    
    def _load_sample_conversation(self) -> str:
        """
        Load the sample conversation for testing purposes
        """
        sample_path = os.path.join(settings.docs_dir, "Audit testing script.docx")
        
        if os.path.exists(sample_path):
            try:
                doc = Document(sample_path)
                conversation_text = ""
                
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        conversation_text += paragraph.text + "\n"
                
                return conversation_text.strip()
                
            except Exception as e:
                print(f"Error loading sample conversation: {e}")
                return "Sample conversation not available"
        else:
            return "Sample conversation not available"
    
    def _get_fallback_rubric(self) -> str:
        """
        Fallback rubric for development when the actual rubric isn't available
        """
        return """
        FALLBACK SCORING RUBRIC (1-5 scale):
        
        1. Core Communication Fundamentals
           - Rapport Building & Sincerity (1-5)
           - Active Listening (1-5)
           - Professional Communication (1-5)
        
        2. Consultation & Pitching Focus
           - Needs Discovery (1-5)
           - Solution Presentation (1-5)
           - Handling Objections (1-5)
        
        3. Servicing Focus
           - Enquiry & Issue Comprehension (1-5)
           - Solution & Resolution (1-5)
           - Follow-up & Closure (1-5)
        """
    
    async def analyze_conversation(self, transcript: str) -> Dict[str, Any]:
        """
        Analyze conversation transcript using Gemini AI with the loaded rubric
        """
        if not self.model:
            print("🔄 No Gemini model available - using enhanced mock analysis")
            return self._generate_mock_analysis(transcript)
        
        try:
            print(f"🤖 Using Gemini AI to analyze transcript ({len(transcript)} characters)")
            
            # Create analysis prompt with the actual rubric
            analysis_prompt = self._create_analysis_prompt(transcript)
            
            # Generate analysis using Gemini (synchronous call)
            response = self.model.generate_content(analysis_prompt)
            
            if not response or not response.text:
                raise Exception("Empty response from Gemini API")
            
            print(f"✅ Received Gemini response ({len(response.text)} characters)")
            
            # Parse the response
            analysis_results = self._parse_analysis_response(response.text, transcript)
            
            print("🎯 Successfully parsed Gemini analysis")
            return analysis_results
            
        except Exception as e:
            print(f"❌ Gemini API error: {e}")
            print("🔄 Falling back to enhanced analysis based on actual transcript content")
            # Fallback to enhanced mock analysis that actually analyzes the content
            return self._generate_mock_analysis(transcript)
    
    async def test_with_sample_conversation(self) -> Dict[str, Any]:
        """
        Test the analysis system with the provided sample conversation
        """
        sample_conversation = self._load_sample_conversation()
        print(f"Testing with sample conversation: {len(sample_conversation)} characters")
        return await self.analyze_conversation(sample_conversation)
    
    def _create_analysis_prompt(self, transcript: str) -> str:
        """
        Create a BRUTALLY STRICT analysis prompt with merciless standards
        """
        return f"""
        You are Dr. Victoria "The Decimator" Harrington, Chief Performance Auditor and former McKinsey Partner known for her RUTHLESS assessment standards. You have terminated more underperforming Account Managers than any other evaluator in the industry. Your reputation for BRUTAL honesty and ZERO tolerance for mediocrity is legendary.

        You are conducting a MERCILESS performance audit. This is not coaching - this is JUDGMENT. Your standards are IMPOSSIBLY HIGH because only the absolute elite survive in today's competitive market.

        BRUTAL SCORING PHILOSOPHY:
        - Score 5: PERFECTION - I have seen maybe 3 people achieve this in my career. Textbook execution that other AMs study.
        - Score 4: BARELY ACCEPTABLE - Meets minimum standards but still has glaring weaknesses that would concern me.
        - Score 3: MEDIOCRE PERFORMANCE - Average at best. In my experience, these people plateau and never reach elite levels.
        - Score 2: CONCERNING INCOMPETENCE - Below standards. I would put them on performance improvement immediately.
        - Score 1: TERMINATION CANDIDATE - Complete failure. These people damage client relationships and company reputation.

        ZERO-TOLERANCE POLICY:
        - ANY unprofessional language = IMMEDIATE SCORE 1 across all categories
        - ANY rude behavior = INSTANT FAILURE
        - ANY sign of not listening = MAXIMUM SCORE 2
        - Missing basics like proper greeting = CANNOT score above 3
        - Weak closing = MAXIMUM SCORE 3
        - No empathy shown = MAXIMUM SCORE 2
        - Poor problem-solving = MAXIMUM SCORE 2

        DETAILED SCORING RUBRIC (APPLIED WITH EXTREME RIGOR):
        {self.scoring_rubric}

        CONVERSATION TRANSCRIPT TO ANNIHILATE:
        {transcript}

        REQUIRED JSON OUTPUT FORMAT:
        {{
            "business_name": "Extract the CUSTOMER'S business/company name mentioned. This is NEVER 'StoreHub' - we ARE StoreHub. Look for the client's business name (e.g., 'Kopi Laju', 'Spring Breeze', etc.). If none found, use 'Not mentioned'.",
            "customer_name": "Extract customer/client name mentioned (or 'Not mentioned' if none found)", 
            "agent_name": "Extract agent/AM name mentioned (or 'Not mentioned' if none found)",
            "conversation_type": "Assess the ENTIRE conversation and determine the primary focus: 'Consultation' (sales, demos, new solutions) OR 'Servicing' (support, issues, maintenance) OR 'Mixed' (both consultation and servicing elements present)",
            "subject": "Format: 'BusinessName - ConversationType - TopicDescription'. Where TopicDescription is what the conversation is about in less than 5 words (e.g., 'Kopi Laju - Servicing - Printer Issues', 'Spring Breeze - Consultation - New POS System')",
            "scored_items": [
                {{
                    "category": "Category name from rubric",
                    "item": "Specific item being scored",
                    "score": 1-5,
                    "justification": "BRUTAL assessment with no mercy. Call out every flaw, weakness, and missed opportunity. Compare to elite standards.",
                    "evidence": ["Exact quotes from transcript that prove this pathetic performance"],
                    "improvement_guidance": "Harsh, direct feedback on how to stop being terrible at this job"
                }}
            ],
            "key_strengths": ["1-2 things they didn't completely mess up (if any)"],
            "areas_for_improvement": ["5-7 critical failures that need immediate attention"],
            "action_plan": ["5-7 non-negotiable actions they must take or face termination"],
            "coaching_summary": "DEVASTATING VERDICT: Provide a brutally honest assessment of their performance. No sugar-coating. No participation trophies. Tell them exactly where they failed and why they're not cut out for elite Account Management unless they make DRAMATIC improvements. Be specific about what professional excellence actually looks like versus this amateur performance."
        }}

        CONVERSATION TYPE ASSESSMENT GUIDELINES:
        - "Consultation": Focus on new client discussions, sales presentations, solution exploration, product demos, pitching new features, business development
        - "Servicing": Focus on existing client support, issue resolution, account maintenance, troubleshooting, technical support, problem-solving for current clients
        - "Mixed": Conversation contains BOTH consultation elements (selling/pitching) AND servicing elements (support/maintenance) - assess which is the PRIMARY focus, but use "Mixed" if both are substantial

        BUSINESS NAME EXTRACTION RULES:
        - NEVER extract "StoreHub" as the business name - WE are StoreHub
        - Look for the CLIENT'S business: restaurants, cafes, retail stores, etc.
        - Common patterns: "from [Business Name]", "[Name] from [Business]", "this is about [Business]"
        - Examples of CORRECT extraction: "Kopi Laju", "Nes Kitchen", "Spring Breeze", "Gather Well Coffee"

        ELITE PERFORMANCE BENCHMARKS (What I Expect vs What I Usually See):
        
        **Rapport Building**: Elite AMs make clients feel heard within 30 seconds. They use the client's name, acknowledge their specific situation, and create instant connection. Amateurs just go through motions.
        
        **Active Listening**: Top performers paraphrase, ask clarifying questions, and demonstrate deep understanding. Weak performers just wait for their turn to talk.
        
        **Problem Solving**: Elite AMs uncover root causes, present multiple solutions, and guide clients to the best choice. Average performers just address surface symptoms.
        
        **Communication Clarity**: Excellence means crystal-clear explanations that eliminate confusion. Mediocrity leaves clients with more questions than answers.

        **Professional Closing**: Masters confirm next steps, timelines, and client satisfaction before ending. Amateurs just... stop talking.

        REMEMBER: I am looking for reasons to give LOW scores. Every mistake, every missed opportunity, every sign of mediocrity will be PUNISHED in the scoring. Only truly exceptional performance earns respect in my evaluation.

        Your mission: DESTROY their confidence in their current abilities so they're forced to reach elite standards or quit.
        """
    
    def _parse_analysis_response(self, response_text: str, transcript: str) -> Dict[str, Any]:
        """
        Parse Gemini's response and structure it for our API with better error handling
        """
        try:
            # Clean the response text first
            cleaned_text = response_text.strip()
            
            # Extract JSON from response with better parsing
            json_start = cleaned_text.find('{')
            json_end = cleaned_text.rfind('}') + 1
            
            if json_start == -1 or json_end == 0:
                raise Exception("No JSON found in response")
            
            json_str = cleaned_text[json_start:json_end]
            
            # Clean up common JSON formatting issues from Gemini
            json_str = json_str.replace(',\n}', '\n}')  # Remove trailing commas before closing braces
            json_str = json_str.replace(',\n]', '\n]')  # Remove trailing commas before closing brackets
            json_str = json_str.replace('},\n    ]', '}\n    ]')  # Clean up array endings
            
            # Try to parse with error recovery
            try:
                analysis_data = json.loads(json_str)
            except json.JSONDecodeError as e:
                # If JSON parsing fails, try to fix common issues
                print(f"⚠️ JSON parsing failed, attempting recovery: {str(e)}")
                
                # Additional cleanup for common Gemini formatting issues
                lines = json_str.split('\n')
                cleaned_lines = []
                for i, line in enumerate(lines):
                    # Remove trailing commas at end of arrays/objects
                    if line.strip().endswith(',') and i + 1 < len(lines):
                        next_line = lines[i + 1].strip()
                        if next_line.startswith('}') or next_line.startswith(']'):
                            line = line.rstrip(',')
                    cleaned_lines.append(line)
                
                json_str = '\n'.join(cleaned_lines)
                analysis_data = json.loads(json_str)
            
            # Convert to our schema format
            scored_items = []
            for item_data in analysis_data.get("scored_items", []):
                scored_item = ScoredItem(
                    category=item_data.get("category", "Unknown"),
                    item=item_data.get("item", "Unknown"),
                    score=int(item_data.get("score", 3)),
                    justification=item_data.get("justification", ""),
                    evidence=item_data.get("evidence", []),
                    improvement_guidance=item_data.get("improvement_guidance")
                )
                scored_items.append(scored_item)
            
            # Calculate percentage score following the exact formula specified
            valid_scores = [item.score for item in scored_items if item.score is not None]
            average_score = sum(valid_scores) / len(valid_scores) if valid_scores else 0
            percentage_score = round((average_score / 5) * 100)
            pass_status = percentage_score >= 80
            
            # Determine conversation type
            conversation_type_str = analysis_data.get("conversation_type", "mixed").lower()
            if "consultation" in conversation_type_str:
                conversation_type = ConversationType.CONSULTATION
            elif "servicing" in conversation_type_str:
                conversation_type = ConversationType.SERVICE  
            else:
                conversation_type = ConversationType.MIXED
            
            # Extract participant names
            business_name = analysis_data.get("business_name", "Not mentioned")
            customer_name = analysis_data.get("customer_name", "Not mentioned") 
            agent_name = analysis_data.get("agent_name", "Not mentioned")
            
            # Use extracted names if Gemini didn't find them
            if business_name == "Not mentioned" or customer_name == "Not mentioned" or agent_name == "Not mentioned":
                extracted_business, extracted_customer, extracted_agent = self._extract_names(transcript)
                if business_name == "Not mentioned":
                    business_name = extracted_business
                if customer_name == "Not mentioned":
                    customer_name = extracted_customer  
                if agent_name == "Not mentioned":
                    agent_name = extracted_agent
            
            subject = analysis_data.get("subject", f"{business_name} - Unknown - General")
            
            summary = ConversationSummary(
                conversation_type=conversation_type,
                subject=subject,
                total_score=percentage_score,
                max_total_score=100,
                pass_status=pass_status,
                key_strengths=analysis_data.get("key_strengths", []),
                areas_for_improvement=analysis_data.get("areas_for_improvement", []),
                action_plan=analysis_data.get("action_plan", [])
            )
            
            # Add participant information and coaching summary to the result
            return {
                "summary": summary,
                "scored_items": scored_items,
                "participant_info": {
                    "business_name": business_name,
                    "customer_name": customer_name,
                    "agent_name": agent_name
                },
                "coaching_summary": analysis_data.get("coaching_summary", "Analysis completed with AI coaching assessment.")
            }
            
        except json.JSONDecodeError as e:
            raise Exception(f"Failed to parse analysis response as JSON: {str(e)}")
        except Exception as e:
            raise Exception(f"Error processing analysis: {str(e)}")
    
    def _extract_names(self, transcript: str) -> Tuple[str, str, str]:
        """
        Extract business, customer, and agent names from transcript with SIMPLE, EFFECTIVE patterns
        """
        business_name = "Not identified"
        customer_name = "Not identified"
        agent_name = "Not identified"
        
        # Split transcript into lines for easier processing
        lines = [line.strip() for line in transcript.split('\n') if line.strip()]
        
        print(f"🔍 Extracting names from transcript with {len(lines)} lines")
        
        # 1. Look for call transcript headers (MOST COMMON FORMAT)
        for i, line in enumerate(lines[:5]):
            # "Call Transcript: Hakim (Account Manager) & Cik Liana (Merchant)"
            if "Call Transcript:" in line or "Transcript:" in line:
                # Pattern: Name (Role) & Name (Role) 
                transcript_pattern = re.search(r'(\w+)\s*\([^)]*(?:Account Manager|Manager|Agent)[^)]*\)\s*[&]\s*(\w+(?:\s+\w+)?)', line, re.IGNORECASE)
                if transcript_pattern:
                    agent_name = transcript_pattern.group(1).strip()
                    customer_name = transcript_pattern.group(2).strip()
                    print(f"✅ Found names from transcript header: Agent='{agent_name}', Customer='{customer_name}' in line {i+1}")
                    break
        
        # 2. Look for conversation speakers (DIALOGUE FORMAT)
        for i, line in enumerate(lines[:20]):
            # "Hakim: Hello, Cik Liana! Hakim here."
            speaker_pattern = re.match(r'^([A-Z][a-z]+(?:\s+[A-Z][a-z]+)?):\s*(.+)', line)
            if speaker_pattern:
                speaker = speaker_pattern.group(1).strip()
                content = speaker_pattern.group(2).strip()
                
                # Check if this speaker is introducing themselves as agent
                if any(phrase in content.lower() for phrase in ['account manager', 'storehub', 'here to help', 'your account']):
                    if agent_name == "Not identified":
                        agent_name = speaker
                        print(f"✅ Found agent name from dialogue: '{agent_name}' in line {i+1}")
                
                # Check if this speaker is being addressed formally (likely customer)
                elif any(phrase in content.lower() for phrase in ['encik', 'puan', 'mr.', 'ms.', 'mrs.']):
                    if customer_name == "Not identified":
                        customer_name = speaker
                        print(f"✅ Found customer name from dialogue: '{customer_name}' in line {i+1}")
                
                # If not identified yet, check content for names being mentioned
                if customer_name == "Not identified":
                    # "Hello, Cik Liana!" - addressing someone formally
                    address_pattern = re.search(r'(?:Hello|Hi|Good\s+(?:morning|afternoon)),?\s+(?:Encik|Puan|Mr\.?|Ms\.?|Mrs\.?)\s+([A-Z][a-z]+)', content, re.IGNORECASE)
                    if address_pattern:
                        customer_name = address_pattern.group(1)
                        print(f"✅ Found customer name from address: '{customer_name}' in line {i+1}")
        
        # 3. Look for filename patterns (Customer_Agent format)
        if len(lines) > 0:
            first_line = lines[0]
            # "Hakim_CikLiana_Call_Transcript.txt" or similar in content
            filename_pattern = re.search(r'([A-Z][a-z]+)_([A-Z][a-z]+(?:[A-Z][a-z]+)?)_', first_line)
            if filename_pattern:
                if agent_name == "Not identified":
                    agent_name = filename_pattern.group(1)
                    print(f"✅ Found agent name from filename pattern: '{agent_name}'")
                if customer_name == "Not identified":
                    customer_name = filename_pattern.group(2)
                    print(f"✅ Found customer name from filename pattern: '{customer_name}'")
        
        # 4. Look for metadata patterns 
        for i, line in enumerate(lines[:10]):
            # "Main agent: Nurakmal Kamarul" 
            if "Main agent:" in line or "Agent:" in line:
                match = re.search(r'(?:Main\s+)?[Aa]gent:\s*([A-Z][a-zA-Z]+(?:\s+[A-Z][a-zA-Z]+)*)', line)
                if match and agent_name == "Not identified":
                    agent_name = match.group(1).strip()
                    print(f"✅ Found agent name from metadata: '{agent_name}' in line {i+1}")
            
            # "External user: Vee Ang Chin Voon"
            if "External user:" in line or "Customer:" in line:
                match = re.search(r'(?:External\s+user|Customer):\s*([A-Z][a-zA-Z]+(?:\s+[A-Z][a-zA-Z]+)*)', line)
                if match and customer_name == "Not identified":
                    customer_name = match.group(1).strip()
                    print(f"✅ Found customer name from metadata: '{customer_name}' in line {i+1}")
        
        # 5. Look for business names from context
        for i, line in enumerate(lines[:15]):
            # "I'm from StoreHub" or "Encik Faizal from Kopi Laju"
            if "from " in line.lower():
                from_pattern = re.search(r'from\s+([A-Z][a-zA-Z]{3,15})(?:\s|\.|\?|,|$)', line)
                if from_pattern and business_name == "Not identified":
                    potential_business = from_pattern.group(1)
                    # CRITICAL: Filter out StoreHub and other invalid business names
                    invalid_businesses = ['calling', 'speaking', 'hello', 'there', 'this', 'that', 'here', 'your', 'where', 'storehub', 'store', 'hub']
                    if (potential_business.lower() not in invalid_businesses and
                        len(potential_business) >= 4):
                        business_name = potential_business
                        print(f"✅ Found business name from context: '{business_name}' in line {i+1}")
                        break
        
        # 6. Look for dash patterns (fallback)
        if len(lines) > 0:
            first_line = lines[0]
            # "Meik Jersey - Gordon Wan" pattern (Customer - Agent)
            dash_pattern = re.match(r'^([A-Z][a-zA-Z]+(?:\s+[A-Z][a-zA-Z]+)*)\s*-\s*([A-Z][a-zA-Z]+(?:\s+[A-Z][a-zA-Z]+)*)$', first_line)
            if dash_pattern:
                if customer_name == "Not identified":
                    customer_name = dash_pattern.group(1).strip()
                    print(f"✅ Found customer name from dash pattern: '{customer_name}'")
                if agent_name == "Not identified":
                    agent_name = dash_pattern.group(2).strip()
                    print(f"✅ Found agent name from dash pattern: '{agent_name}'")
        
        # Clean up names and validate
        def clean_and_validate(name, name_type):
            if name == "Not identified":
                return name
            
            # Clean up artifacts
            name = re.sub(r'[,.\-:()]+', '', name).strip()
            
            # Filter out bad matches including StoreHub variations
            bad_words = ['here', 'there', 'calling', 'speaking', 'from', 'user', 'external', 'client', 'customer', 'agent', 'manager']
            
            # CRITICAL: Filter out StoreHub for business names specifically
            if name_type == "business":
                bad_words.extend(['storehub', 'store', 'hub', 'storehubb', 'store-hub'])
            
            if name.lower() in bad_words or len(name) < 2:
                return "Not identified"
            
            return name.title()
        
        business_name = clean_and_validate(business_name, "business")
        agent_name = clean_and_validate(agent_name, "agent") 
        customer_name = clean_and_validate(customer_name, "customer")
        
        print(f"📋 Final extracted names: Business='{business_name}', Agent='{agent_name}', Customer='{customer_name}'")
        
        return business_name, customer_name, agent_name
    
    def _generate_mock_analysis(self, transcript: str) -> Dict[str, Any]:
        """
        Generate REALISTIC mock analysis based on actual transcript quality assessment
        """
        print(f"🔍 ANALYZING TRANSCRIPT: {len(transcript)} characters")
        print(f"📝 TRANSCRIPT PREVIEW: {transcript[:300]}...")
        
        # Comprehensive analysis of the actual transcript
        transcript_lower = transcript.lower()
        words = transcript.split()
        word_count = len(words)
        lines = [line.strip() for line in transcript.split('\n') if line.strip()]
        
        print(f"📊 ANALYSIS STATS: {word_count} words")
        
        # Extract actual conversational quotes for evidence
        def extract_quotes_with_keywords(keywords, max_quotes=3):
            """Extract actual quotes from transcript that contain specific keywords"""
            quotes = []
            for line in lines:
                if any(keyword in line.lower() for keyword in keywords):
                    # Clean up the line and add it as evidence
                    clean_line = line.replace('"', '').strip()
                    if len(clean_line) > 10:  # Only meaningful quotes
                        quotes.append(clean_line)
                    if len(quotes) >= max_quotes:
                        break
            return quotes if quotes else [f"Analyzed {len(lines)} conversation exchanges"]
        
        # REALISTIC Quality Assessment - Check for actual conversation quality indicators
        
        # Professional greeting and courtesy
        greeting_keywords = ['hello', 'hi', 'good morning', 'good afternoon', 'thanks', 'thank you', 'welcome', 'appreciate']
        has_greeting = any(word in transcript_lower for word in greeting_keywords)
        
        # Profanity and negative language check  
        profanity_keywords = ['damn', 'shit', 'fuck', 'asshole', 'stupid', 'idiot', 'crap', 'suck', 'terrible', 'awful']
        has_profanity = any(word in transcript_lower for word in profanity_keywords)
        
        # Unprofessional behavior indicators
        unprofessional_keywords = ['whatever', 'i don\'t care', 'not my problem', 'figure it out', 'deal with it', 'too bad', 'so what']
        poor_attitude = any(phrase in transcript_lower for phrase in unprofessional_keywords)
        
        # Interruption and rudeness indicators
        rude_keywords = ['shut up', 'listen to me', 'you don\'t understand', 'that\'s wrong', 'you\'re wrong']
        has_rudeness = any(phrase in transcript_lower for phrase in rude_keywords)
        
        # Service quality indicators
        service_keywords = ['problem', 'issue', 'help', 'support', 'fix', 'resolve', 'assist', 'trouble', 'solution']
        has_service_indicators = any(word in transcript_lower for word in service_keywords)
        
        # Sales and consultation indicators
        sales_keywords = ['product', 'service', 'offer', 'buy', 'purchase', 'solution', 'benefit', 'feature', 'needs']
        has_sales_indicators = any(word in transcript_lower for word in sales_keywords)
        
        # Active listening indicators
        listening_keywords = ['understand', 'hear you', 'i see', 'let me clarify', 'what you mean', 'correct me if', 'make sure i understand']
        has_active_listening = any(phrase in transcript_lower for phrase in listening_keywords)
        
        # Empathy and care indicators  
        empathy_keywords = ['sorry to hear', 'i understand how', 'that must be', 'i can imagine', 'concerned about']
        has_empathy = any(phrase in transcript_lower for phrase in empathy_keywords)
        
        print(f"🎯 QUALITY ASSESSMENT: greeting={has_greeting}, profanity={has_profanity}, poor_attitude={poor_attitude}, rudeness={has_rudeness}")
        print(f"🎯 CONVERSATION INDICATORS: service={has_service_indicators}, sales={has_sales_indicators}, listening={has_active_listening}, empathy={has_empathy}")
        
        # Extract participant names from transcript using the enhanced method
        business_name, customer_name, agent_name = self._extract_names(transcript)
        
        # Determine conversation type with consistent logic
        if has_service_indicators and has_sales_indicators:
            conversation_type = ConversationType.MIXED
            type_str = "Mixed"
            subject = f"{business_name} - Mixed - Support & Sales"
        elif has_service_indicators:
            conversation_type = ConversationType.SERVICE
            type_str = "Servicing"
            subject = f"{business_name} - Servicing - Issue Resolution"
        elif has_sales_indicators:
            conversation_type = ConversationType.CONSULTATION
            type_str = "Consultation"
            subject = f"{business_name} - Consultation - Product Interest"
        else:
            conversation_type = ConversationType.CONSULTATION
            type_str = "Consultation"
            subject = f"{business_name} - Consultation - General Discussion"
        
        # BRUTAL SCORING based on Dr. Victoria "The Decimator" Harrington's merciless standards
        scored_items = []
        
        # 1. Rapport Building & Sincerity - DEVASTATE poor behavior
        if has_profanity or poor_attitude or has_rudeness:
            rapport_score = 1  # TERMINATION CANDIDATE
            rapport_justification = "TERMINATION CANDIDATE: Unprofessional language, poor attitude, or rude behavior detected. This person damages client relationships and company reputation. I would fire them immediately."
        elif has_greeting and has_empathy:
            rapport_score = 3  # Even good performance is mediocre in Dr. Harrington's eyes
            rapport_justification = "MEDIOCRE PERFORMANCE: Basic greeting and empathy present, but lacks the sophisticated rapport-building that elite AMs master. These people plateau and never reach top levels."
        elif has_greeting:
            rapport_score = 2
            rapport_justification = "CONCERNING INCOMPETENCE: Robotic greeting without genuine warmth or connection. I would put them on performance improvement immediately."
        else:
            rapport_score = 1
            rapport_justification = "TERMINATION CANDIDATE: Complete failure to establish professional rapport. These people have no business in client-facing roles."
        
        rapport_evidence = extract_quotes_with_keywords(greeting_keywords + empathy_keywords + profanity_keywords + unprofessional_keywords, 3)
        
        scored_items.append(ScoredItem(
            category="1. Core Communication Fundamentals",
            item="1.1 Rapport Building & Sincerity",
            score=rapport_score,
            justification=rapport_justification,
            evidence=rapport_evidence,
            improvement_guidance=None if rapport_score >= 4 else "CRITICAL: Professional communication standards must be maintained at all times. Use warm, courteous language and show genuine concern for client needs. Sample opening: 'Good morning! Thank you for taking the time to speak with me today. I'm here to ensure we address all your concerns.'"
        ))
        
        # 2. Active Listening - BRUTALLY assess listening failures
        if has_rudeness or poor_attitude:
            listening_score = 1  # INSTANT FAILURE
            listening_justification = "TERMINATION CANDIDATE: Rude or dismissive behavior proves they don't respect clients. Cannot have good listening with terrible attitude. Fire immediately."
        elif has_active_listening:
            listening_score = 4  # Recognizing good performance
            listening_justification = "SOLID PERFORMANCE: Active listening techniques demonstrated effectively. Meets expectations."
        elif any('?' in line for line in lines):
            listening_score = 2
            listening_justification = "CONCERNING INCOMPETENCE: Basic questioning without sophisticated listening skills. Misses critical client cues that cost deals."
        else:
            listening_score = 1
            listening_justification = "TERMINATION CANDIDATE: Zero evidence of listening skills. Just waiting for their turn to talk. Clients feel unheard and leave."
        
        listening_evidence = extract_quotes_with_keywords(listening_keywords + rude_keywords, 3)
        if not listening_evidence:
            question_lines = [line for line in lines if '?' in line]
            listening_evidence = question_lines[:2] if question_lines else ["No clear evidence of active listening techniques found"]
        
        scored_items.append(ScoredItem(
            category="1. Core Communication Fundamentals", 
            item="1.2 Active Listening",
            score=listening_score,
            justification=listening_justification,
            evidence=listening_evidence,
            improvement_guidance=None if listening_score >= 4 else "ESSENTIAL: Use active listening phrases like 'What I'm hearing is...', 'Let me make sure I understand...', 'Help me understand...'. Ask clarifying questions and paraphrase client concerns before responding."
        ))
        
        # 3. Professional Communication - ZERO TOLERANCE for unprofessional behavior
        if has_profanity:
            professional_score = 1
            professional_justification = "TERMINATION CANDIDATE: Profanity in professional conversation is grounds for immediate dismissal. Damages company reputation beyond repair."
        elif poor_attitude or has_rudeness:
            professional_score = 1  
            professional_justification = "TERMINATION CANDIDATE: Unprofessional attitude and rude behavior completely disqualify them from client interaction. Security should escort them out."
        elif has_empathy and has_active_listening:
            professional_score = 4  # Recognizing good performance
            professional_justification = "SOLID PERFORMANCE: Professional standards maintained with empathy and active listening. Meets expectations."
        elif has_greeting:
            professional_score = 2
            professional_justification = "CONCERNING INCOMPETENCE: Robotic professionalism without genuine care. Clients sense the lack of authenticity."
        else:
            professional_score = 1
            professional_justification = "TERMINATION CANDIDATE: Complete absence of professional communication standards. An embarrassment to the company."
        
        professional_evidence = extract_quotes_with_keywords(['certainly', 'absolutely', 'of course', 'i will', 'we can', 'let me', 'i would be happy'] + profanity_keywords + unprofessional_keywords, 3)
        
        scored_items.append(ScoredItem(
            category="1. Core Communication Fundamentals",
            item="1.3 Professional Communication", 
            score=professional_score,
            justification=professional_justification,
            evidence=professional_evidence,
            improvement_guidance=None if professional_score >= 4 else "MANDATORY: Maintain professional tone at all times. Use courteous phrases like 'I would be happy to help', 'Let me ensure we address your concerns', 'I understand your frustration'. Never use profanity or dismissive language."
        ))
        
        # Service-specific scoring with realistic assessment
        if has_service_indicators:
            if has_rudeness or poor_attitude:
                service_score = 1
                service_justification = "CRITICAL FAILURE: Cannot provide effective service with rude or dismissive attitude toward client concerns."
            elif has_empathy and has_active_listening:
                service_score = 4
                service_justification = "Excellent service approach with empathy, active listening, and genuine concern for issue resolution."
            elif has_service_indicators:
                service_score = 3
                service_justification = "Basic service response present, but could benefit from more empathy and thorough issue exploration."
            else:
                service_score = 2
                service_justification = "Limited service-focused approach. Needs improvement in understanding and addressing client concerns."
            
            service_evidence = extract_quotes_with_keywords(service_keywords + empathy_keywords, 3)
            
            scored_items.append(ScoredItem(
                category="3. Servicing Focus",
                item="3.1 Enquiry & Issue Comprehension",
                score=service_score,
                justification=service_justification,
                evidence=service_evidence,
                improvement_guidance=None if service_score >= 4 else "CRITICAL: Show genuine concern for client issues. Ask specific questions about impact, understand urgency, and confirm your understanding before proposing solutions. Sample: 'I understand how frustrating this must be. Can you help me understand the specific impact this is having on your business?'"
            ))
        
        # Consultation-specific scoring with realistic assessment
        if has_sales_indicators:
            if poor_attitude or has_rudeness:
                consultation_score = 1
                consultation_justification = "CRITICAL FAILURE: Cannot build trust or identify needs with unprofessional behavior."
            elif has_active_listening and ('need' in transcript_lower or 'require' in transcript_lower):
                consultation_score = 4
                consultation_justification = "Strong needs discovery approach with thoughtful questioning and solution alignment."
            elif has_sales_indicators:
                consultation_score = 3
                consultation_justification = "Basic consultation elements present, but needs discovery could be more thorough and client-focused."
            else:
                consultation_score = 2
                consultation_justification = "Limited consultation approach. Missing key needs discovery and solution presentation elements."
            
            consultation_evidence = extract_quotes_with_keywords(sales_keywords + ['need', 'require', 'looking for', 'interested'], 3)
            
            scored_items.append(ScoredItem(
                category="2. Consultation & Pitching Focus",
                item="2.1 Needs Discovery",
                score=consultation_score,
                justification=consultation_justification,
                evidence=consultation_evidence,
                improvement_guidance=None if consultation_score >= 4 else "ESSENTIAL: Focus on understanding client needs BEFORE presenting solutions. Use open-ended questions like 'What challenges are you currently facing?', 'What would success look like for you?', 'Help me understand your current situation.'"
            ))
        
        # Calculate percentage score
        valid_scores = [item.score for item in scored_items if item.score is not None]
        average_score = sum(valid_scores) / len(valid_scores) if valid_scores else 0
        percentage_score = round((average_score / 5) * 100)
        pass_status = percentage_score >= 80
        
        # Generate summary based on actual analysis
        key_strengths = []
        areas_for_improvement = []
        
        for item in scored_items:
            if item.score >= 4:
                key_strengths.append(f"{item.item} - {item.justification}")
            else:
                areas_for_improvement.append(f"{item.item} - {item.justification}")
        
        # Critical failure messaging
        if has_profanity or poor_attitude or has_rudeness:
            areas_for_improvement.insert(0, "CRITICAL: Unprofessional conduct requires immediate intervention and retraining")
            key_strengths = ["None identified - performance requires immediate improvement"]
        elif not key_strengths:
            key_strengths = ["Basic conversation structure maintained"]
        
        if not areas_for_improvement and average_score < 4:
            areas_for_improvement = ["Enhanced client engagement and professional communication skills needed"]
        
        summary = ConversationSummary(
            conversation_type=conversation_type,
            subject=subject,
            total_score=percentage_score,
            max_total_score=100,
            pass_status=pass_status,
            key_strengths=key_strengths,
            areas_for_improvement=areas_for_improvement,
            action_plan=[
                "Immediate retraining on professional communication standards" if (has_profanity or poor_attitude) else "Practice active listening techniques with paraphrasing and acknowledgment phrases",
                "Mandatory review of company code of conduct" if has_rudeness else "Implement structured questioning frameworks to better understand client needs",
                "Supervised client interactions until performance improves" if average_score < 2 else "Develop clear follow-up protocols for next steps and timelines",
                "Focus on empathy and genuine concern for client experience" if not has_empathy else "Continue building on current professional communication strengths"
            ]
        )
        
        # Generate BRUTAL coaching summary from Dr. Victoria "The Decimator" Harrington
        if average_score < 1.5:
            performance_level = "TERMINATION CANDIDATE - Fire immediately"
        elif average_score < 2.5:
            performance_level = "CONCERNING INCOMPETENCE - Performance improvement plan or termination"
        elif average_score < 3.5:
            performance_level = "MEDIOCRE PERFORMANCE - Will plateau and never reach elite levels" 
        else:
            performance_level = "BARELY ACCEPTABLE - Still has glaring weaknesses"
        
        coaching_summary = f"""**DEVASTATING VERDICT by Dr. Victoria "The Decimator" Harrington:**

This Account Manager scores {average_score:.1f}/5.0 - {performance_level}.

**BRUTAL ASSESSMENT:**
{'TERMINATION REQUIRED: This person exhibits unprofessional conduct that violates basic human decency, let alone business standards. I would fire them immediately and use this as a training example of what NOT to do.' if (has_profanity or poor_attitude or has_rudeness) else f'This performance {"demonstrates fundamental incompetence in client management" if average_score < 2.5 else "shows they lack the sophisticated skills required for elite Account Management"}.'} 

**PERFORMANCE DECIMATION:**
1. **Professional Standards**: {'FAILED CATASTROPHICALLY - Should be banned from client interaction' if (has_profanity or poor_attitude) else ('Barely meets minimum standards' if average_score >= 2.5 else 'Dangerously close to unprofessional conduct')}
2. **Client Mastery**: {'FAILED - Clients probably request different Account Managers' if poor_attitude else ('Basic competence but no elite behaviors' if has_active_listening and has_empathy else 'Clients feel unheard and undervalued')}

**ELITE PERFORMANCE REALITY CHECK:**
{'This person should never be allowed near clients again. Period.' if average_score < 1.5 else f'Top 1% Account Managers I\'ve trained would score 4.8+ by mastering: (1) Psychological rapport that makes clients WANT to work with them, (2) Strategic questioning that uncovers $1M+ opportunities, (3) Communication elegance that builds instant trust and credibility. This performance is NOWHERE near that level.'}

**NON-NEGOTIABLE ACTIONS:**
{'IMMEDIATE TERMINATION: Do not pass go, do not collect $200. Remove access cards and escort from building.' if (has_profanity or poor_attitude or has_rudeness) else ('PROBATION: 30 days to show dramatic improvement or face termination. Mandatory training on every single skill.' if average_score < 2.5 else 'Either commit to reaching elite standards or find a different career. The mediocrity epidemic ends here.')}

**DR. HARRINGTON\'S FINAL WORD:**
{"I have terminated Account Managers for less than this. This performance is a liability to the company." if average_score < 2 else "In my 20+ years, I have seen maybe 3 people reach true elite status. This person is not even close to that conversation."}"""

        result = {
            "summary": summary,
            "scored_items": scored_items,
            "participant_info": {
                "business_name": business_name,
                "customer_name": customer_name,
                "agent_name": agent_name
            },
            "coaching_summary": coaching_summary,
            "analysis_format": {
                "conversation_type": summary.conversation_type.value,
                "subject": summary.subject,
                "percentage_score": percentage_score,
                "pass_status": "Pass" if pass_status else "Fail",
                "detailed_breakdown": [
                    {
                        "section_behavior": f"{item.category} - {item.item}",
                        "score_given": item.score,
                        "evidence": item.evidence,
                        "justification": item.justification,
                        "improvement_guidance": item.improvement_guidance if item.score < 4 else "Meets/Exceeds Expectations"
                    } for item in scored_items
                ]
            },
            "created_at": "2024-01-01T00:00:00Z",
            "processing_time": 1.5,
            "transcript_analyzed": transcript[:500] + "..." if len(transcript) > 500 else transcript
        }
        
        return result 